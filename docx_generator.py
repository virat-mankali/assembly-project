import tempfile

from docx import Document
from docx.shared import Inches, Pt

HEADER_MARKER = "mcm/nns"
SPEAKER_PREFIXES = ("} Sri ", "} MR. SPEAKER", "} THE SPEAKER", "} MADAM SPEAKER")


def _apply_paragraph_font(paragraph, *, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold


def _apply_paragraph_spacing(paragraph, *, space_after: int = 6) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(space_after)


def _is_speaker_label(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith(SPEAKER_PREFIXES) and stripped.endswith(":")


def _has_session_header(lines: list[str]) -> bool:
    return any(HEADER_MARKER in line for line in lines[:4])


def _ensure_session_header(formatted_text: str) -> list[str]:
    lines = formatted_text.splitlines()
    if _has_session_header(lines):
        return lines

    return ["", f"                {HEADER_MARKER}", "", *lines]


def generate_docx(formatted_text: str, session_label: str = "") -> str:
    if not formatted_text.strip():
        raise ValueError("Formatted text is empty")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    lines = _ensure_session_header(formatted_text)
    last_paragraph = None

    for index, line in enumerate(lines):
        clean_line = line.strip()
        if not clean_line:
            if index == 0 and _has_session_header(lines):
                paragraph = doc.add_paragraph("")
                _apply_paragraph_spacing(paragraph)
                last_paragraph = paragraph
                continue
            if last_paragraph is not None:
                last_paragraph.paragraph_format.space_after = Pt(12)
            continue

        is_speaker_label = _is_speaker_label(clean_line)
        paragraph_text = line.rstrip() if clean_line == HEADER_MARKER else clean_line
        paragraph = doc.add_paragraph(paragraph_text)
        _apply_paragraph_font(paragraph, bold=is_speaker_label)
        _apply_paragraph_spacing(paragraph)

        if line.startswith((" ", "\t")) and not is_speaker_label and clean_line != HEADER_MARKER:
            paragraph.paragraph_format.left_indent = Inches(0.25)

        last_paragraph = paragraph

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    return tmp.name
