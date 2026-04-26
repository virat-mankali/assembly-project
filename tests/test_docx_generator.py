import os
import unittest

from docx import Document
from docx.shared import Inches, Pt

from docx_generator import generate_docx


class DocxGeneratorTests(unittest.TestCase):
    def test_applies_legislative_docx_formatting(self) -> None:
        path = generate_docx(
            "\n".join(
                [
                    "26.03.2026          11.40 P.M.",
                    "                mcm/nns",
                    "",
                    "} Sri K.T. Rama Rao (BRS):",
                    "    ఇది ఒక పరీక్ష డైలాగ్.",
                    "",
                    "} Sri D. Sridhar Babu (Minister):",
                    "    మరొక డైలాగ్.",
                ]
            )
        )

        try:
            doc = Document(path)
            section = doc.sections[0]
            self.assertEqual(section.left_margin, Inches(1.25))
            self.assertEqual(section.top_margin, Inches(1))
            self.assertEqual(section.right_margin, Inches(1))
            self.assertEqual(section.bottom_margin, Inches(1))

            paragraphs = doc.paragraphs
            self.assertEqual(paragraphs[0].text, "26.03.2026          11.40 P.M.")
            self.assertEqual(paragraphs[1].text, "                mcm/nns")
            self.assertEqual(paragraphs[2].text, "} Sri K.T. Rama Rao (BRS):")
            self.assertTrue(paragraphs[2].runs[0].bold)
            self.assertFalse(paragraphs[3].runs[0].bold)
            self.assertEqual(paragraphs[3].paragraph_format.left_indent, Inches(0.25))

            for paragraph in paragraphs:
                self.assertEqual(paragraph.paragraph_format.line_spacing, 1.5)
                for run in paragraph.runs:
                    self.assertEqual(run.font.name, "Times New Roman")
                    self.assertEqual(run.font.size, Pt(12))

            self.assertEqual(paragraphs[3].paragraph_format.space_after, Pt(12))
            self.assertEqual(paragraphs[5].paragraph_format.space_after, Pt(6))
        finally:
            os.unlink(path)

    def test_adds_blank_session_header_when_missing(self) -> None:
        path = generate_docx(
            "\n".join(
                [
                    "} Sri K.T. Rama Rao (BRS):",
                    "    Dialogue.",
                ]
            )
        )

        try:
            doc = Document(path)
            self.assertEqual(doc.paragraphs[0].text, "")
            self.assertEqual(doc.paragraphs[1].text, "                mcm/nns")
            self.assertEqual(doc.paragraphs[2].text, "} Sri K.T. Rama Rao (BRS):")
        finally:
            os.unlink(path)
